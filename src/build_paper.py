"""
T8 — build paper.docx in the IEEE conference format.

The supplied IEEE_Template.docx is an ISO-Strict OOXML package: its relationship
types live under purl.oclc.org rather than schemas.openxmlformats.org, which
python-docx does not recognise. A namespace-rewritten copy is produced first,
purely so the template's own styles, page setup and column geometry are
inherited rather than reinvented.

Layout produced:
    section 1   one column    title and author block
    section 2   two columns   abstract onwards
    section 3   one column    Fig. 4, which spans both columns
    section 4   two columns   remainder, ending with the references

Run:  python src/build_paper.py        Writes paper.docx
"""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "src"))

import paper_content as C  # noqa: E402

TEMPLATE = ROOT / "IEEE_Template.docx"
OUT = ROOT / "paper.docx"
FIGS = ROOT / "figures"
TMP = ROOT / ".paper_template_tmp.docx"

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

STRICT = "http://purl.oclc.org/ooxml"
NS_FIX = {
    f"{STRICT}/officeDocument/relationships":
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    f"{STRICT}/wordprocessingml/main":
        "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    f"{STRICT}/drawingml/main":
        "http://schemas.openxmlformats.org/drawingml/2006/main",
}

COL_W = Inches(3.30)
FULL_W = Inches(6.90)


# =============================================================================
# template plumbing
# =============================================================================
def transitional_copy(dst: Path) -> Path:
    with zipfile.ZipFile(TEMPLATE) as src, \
            zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as out:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", "replace")
                for a, b in NS_FIX.items():
                    text = text.replace(a, b)
                data = text.encode("utf-8")
            out.writestr(item, data)
    return dst


def clear_body(doc) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_columns(section, n: int, space_twips: int = 360) -> None:
    sect = section._sectPr
    for old in sect.xpath("./w:cols"):
        sect.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sect.append(cols)


def copy_page_setup(src, dst) -> None:
    for attr in ("orientation", "page_width", "page_height", "left_margin",
                 "right_margin", "top_margin", "bottom_margin",
                 "header_distance", "footer_distance"):
        try:
            setattr(dst, attr, getattr(src, attr))
        except Exception:
            pass


def new_section(doc, cols: int):
    prev = doc.sections[-1]
    s = doc.add_section(WD_SECTION.CONTINUOUS)
    copy_page_setup(prev, s)
    set_columns(s, cols)
    return s


# =============================================================================
# content helpers
# =============================================================================
def body(doc, text, indent=True, style="Body Text"):
    p = doc.add_paragraph(style=style)
    p.add_run(CITE.resolve(text))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.16)
    p.paragraph_format.space_after = Pt(2)
    return p


def heading(doc, text, level=1):
    return doc.add_paragraph(CITE.resolve(text), style=f"Heading {level}")


def equation(doc, text, number):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(1.55), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.RIGHT)
    pf.space_before, pf.space_after = Pt(4), Pt(4)
    r = p.add_run("\t" + text + "\t(" + str(number) + ")")
    r.italic = True
    r.font.size = Pt(9)
    return p


def figure(doc, name, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIGS / name), width=width)
    c = doc.add_paragraph(style="figure caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(CITE.resolve(caption))
    c.paragraph_format.space_after = Pt(8)


def set_table_borders(table) -> None:
    """The template defines no table style, so the rules are drawn here.

    IEEE conference tables are ruled, not boxed in every cell: a line above the
    header, one below it, and one under the last row.
    """
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in ("top", "bottom"):
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
        elif edge == "insideH":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    pr.append(borders)


def table_block(doc, number, title, header, rows, widths=None, note=None,
                font_pt=6.5, header_pt=6.5):
    cap = doc.add_paragraph(style="table head")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"TABLE {number}").bold = False
    cap.add_run("\n" + title.upper())
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(2)

    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = False
    set_table_borders(t)
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_after = Pt(0)
        r = pp.add_run(CITE.resolve(h))
        r.bold = True
        r.font.size = Pt(header_pt)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            pp = cells[i].paragraphs[0]
            pp.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                            else WD_ALIGN_PARAGRAPH.CENTER)
            pp.paragraph_format.space_after = Pt(0)
            r = pp.add_run(CITE.resolve(str(v)))
            r.font.size = Pt(font_pt)
            if str(v).startswith("This article"):
                r.bold = True
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    if note:
        n = doc.add_paragraph(style="table footnote")
        run = n.add_run(CITE.resolve(note))
        run.font.size = Pt(6)
        n.paragraph_format.space_after = Pt(8)
    return t


# =============================================================================
# references
# =============================================================================
class Citations:
    """Assigns reference numbers in order of first appearance.

    The prose carries symbolic keys; numbers are handed out the first time a
    key is written into the document. Because the document is assembled
    strictly front to back, assignment order is appearance order, which is
    what IEEE requires. Renumbering after an edit is therefore automatic
    rather than a manual pass that can silently go wrong.
    """

    KEY = re.compile(r"\{([a-z][A-Za-z0-9]*)\}")

    def __init__(self) -> None:
        self.order: list[str] = []

    def resolve(self, text: str) -> str:
        def one(m):
            key = m.group(1)
            if key not in self.order:
                self.order.append(key)
            return f"[{self.order.index(key) + 1}]"
        return self.KEY.sub(one, text)


CITE = Citations()


def load_bib():
    t = (ROOT / "paper" / "references.bib").read_text(encoding="utf-8")
    out = {}
    for m in re.finditer(r"@(\w+)\{([^,]+),(.*?)\n\}", t, re.S):
        kind, key, blk = m.group(1), m.group(2).strip(), m.group(3)
        f = {k.lower(): re.sub(r"\s+", " ", v).strip()
             for k, v in re.findall(r"(\w+)\s*=\s*\{(.*?)\}(?=,\s*\n|\s*$)", blk, re.S)}
        f["_kind"] = kind
        out[key] = f
    return out


def clean(s: str) -> str:
    s = re.sub(r"\{\\[\'\"^`~v]\{?(\w)\}?\}", r"\1", s)
    s = re.sub(r"\\[\'\"^`~v]\{?(\w)\}?", r"\1", s)
    return (s.replace("{", "").replace("}", "")
             .replace("\\&", "&").replace("---", "\u2013").replace("--", "\u2013"))


def authors_ieee(raw: str) -> str:
    raw = clean(raw)
    if raw.startswith("Ember") or "Organization for Standard" in raw:
        return raw
    out = []
    for n in [x.strip() for x in raw.split(" and ")]:
        if "," in n:
            last, first = [x.strip() for x in n.split(",", 1)]
        else:
            bits = n.split()
            last, first = bits[-1], " ".join(bits[:-1])
        init = " ".join(f"{p[0]}." for p in first.replace("-", " ").split() if p)
        out.append(f"{init} {last}".strip())
    if len(out) > 6:
        return out[0] + " et al."
    if len(out) == 1:
        return out[0]
    return ", ".join(out[:-1]) + ", and " + out[-1]


def reference_text(f: dict) -> str:
    a, kind = authors_ieee(f.get("author", "")), f["_kind"]
    title = clean(f.get("title", ""))
    venue = clean(f.get("journal") or f.get("booktitle") or f.get("howpublished") or "")
    s = f"{a}, \u201c{title},\u201d "
    if kind == "article":
        s += venue
        if f.get("volume"):
            s += f", vol. {f['volume']}"
        if f.get("number"):
            s += f", no. {f['number']}"
        if f.get("pages"):
            s += f", pp. {clean(f['pages'])}"
        s += f", {f.get('year','')}."
    elif kind in ("inproceedings", "incollection"):
        s += f"in {venue}"
        if f.get("volume"):
            s += f", vol. {f['volume']}"
        if f.get("pages"):
            s += f", pp. {clean(f['pages'])}"
        s += f", {f.get('year','')}."
    elif kind == "techreport":
        s += f"{clean(f.get('institution',''))}, {f.get('year','')}."
    else:
        s += f"{venue}, {f.get('year','')}."
    if f.get("doi"):
        s += f" doi: {f['doi']}."
    return re.sub(r"\s{2,}", " ", s).replace(" ,", ",")


# =============================================================================
# assembly
# =============================================================================
def build() -> Document:
    doc = Document(str(transitional_copy(TMP)))
    clear_body(doc)

    s0 = doc.sections[0]
    set_columns(s0, 1)

    # --- title block ------------------------------------------------------
    t = doc.add_paragraph(C.TITLE, style="paper title")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a = doc.add_paragraph(style="Author")
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(C.AUTHOR_LINES):
        if i:
            a.add_run("\n")
        r = a.add_run(line)
        if i == 0:
            r.bold = True
    a.paragraph_format.space_after = Pt(10)

    # --- two-column body --------------------------------------------------
    new_section(doc, 2)

    ab = doc.add_paragraph(style="Abstract")
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ab.add_run("Abstract\u2014").bold = True
    ab.add_run(CITE.resolve(C.ABSTRACT)).bold = True
    kw = doc.add_paragraph(style="Keywords")
    kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw.add_run("Keywords\u2014").bold = True
    kw.add_run(C.KEYWORDS)
    kw.paragraph_format.space_after = Pt(6)

    # --- I. Introduction --------------------------------------------------
    heading(doc, "Introduction")
    for i, p in enumerate(C.INTRO):
        body(doc, p, indent=bool(i))
    body(doc, "The contributions of this article are threefold:")
    for i, c in enumerate(C.CONTRIBUTIONS):
        p = body(doc, f"({'i'*(i+1) if i < 3 else i+1}) {c}", indent=False)
        p.paragraph_format.left_indent = Inches(0.10)
    body(doc, C.INTRO_STRUCTURE)

    # --- II. Related Work -------------------------------------------------
    heading(doc, "Related Work")
    for title, paras in C.RELATED:
        heading(doc, title, 2)
        for i, p in enumerate(paras):
            body(doc, p, indent=bool(i))
    table_block(doc, "I", "Positioning against representative studies",
                C.TABLE1_HEADER, C.TABLE1_ROWS,
                widths=[1.02, 0.46, 0.60, 0.58, 0.60], note=C.TABLE1_NOTE)
    body(doc, C.RELATED_GAP, indent=False)

    # --- III. Proposed Framework -----------------------------------------
    heading(doc, "Proposed Framework")
    heading(doc, "Scope and Assumptions", 2)
    body(doc, C.FW_SCOPE, indent=False)

    heading(doc, "Architecture", 2)
    body(doc, C.FW_ARCH, indent=False)
    figure(doc, "fig1_architecture.png", COL_W, C.FIG1_CAPTION)

    heading(doc, "Machine Digital Twin", 2)
    body(doc, C.FW_MDT[0], indent=False)
    equation(doc, C.EQUATIONS["health"], 1)
    body(doc, C.FW_MDT[1], indent=False)

    heading(doc, "Human Digital Twin", 2)
    body(doc, C.FW_HDT[0], indent=False)
    equation(doc, C.EQUATIONS["fatigue"], 2)
    body(doc, C.FW_HDT[1], indent=False)
    body(doc, C.FW_HDT[2])
    equation(doc, C.EQUATIONS["fhat"], 3)
    body(doc, C.FW_HDT[3], indent=False)
    body(doc, C.FW_HDT[4])
    equation(doc, C.EQUATIONS["rula"], 4)
    body(doc, C.FW_HDT[5], indent=False)

    heading(doc, "Twin Coupling", 2)
    body(doc, C.FW_COUPLING[0], indent=False)
    equation(doc, C.EQUATIONS["quality"], 5)
    body(doc, C.FW_COUPLING[1], indent=False)
    figure(doc, "fig2_dataflow.png", COL_W, C.FIG2_CAPTION)

    heading(doc, "Objective and Constraints", 2)
    body(doc, C.FW_OBJ[0], indent=False)
    equation(doc, C.EQUATIONS["objective"], 6)
    body(doc, C.FW_OBJ[1], indent=False)
    body(doc, C.FW_OBJ[2])
    body(doc, C.FW_OBJ[3])

    # --- IV. Implementation ----------------------------------------------
    heading(doc, "Implementation and Experimental Setup")
    for j, (title, paras) in enumerate(C.IMPL):
        heading(doc, title, 2)
        for i, p in enumerate(paras):
            body(doc, p, indent=bool(i))
        if j == 0:
            table_block(doc, "II", "Simulation parameters and their sources",
                        C.TABLE2_HEADER, C.TABLE2_ROWS,
                        widths=[1.35, 1.05, 0.90])

    # --- V. Results -------------------------------------------------------
    heading(doc, "Results and Discussion")
    for j, (title, paras) in enumerate(C.RESULTS):
        heading(doc, title, 2)
        for i, p in enumerate(paras):
            body(doc, p, indent=bool(i))
        if j == 0:
            table_block(doc, "III", "KPI comparison under high demand (S2)",
                        *kpi_table_data(), widths=[0.78, 0.60, 0.60, 0.60, 0.42, 0.30],
                        note=C.TABLE3_NOTE)
            figure(doc, "fig3_tradeoff.png", COL_W, C.FIG3_CAPTION)

    # --- Fig. 4 spans both columns ---------------------------------------
    new_section(doc, 1)
    figure(doc, "fig4_comparison.png", FULL_W, C.FIG4_CAPTION)
    new_section(doc, 2)

    # --- VI, VII ----------------------------------------------------------
    heading(doc, "Limitations and Future Work")
    for i, p in enumerate(C.LIMITATIONS):
        body(doc, p, indent=bool(i))

    heading(doc, "Conclusion")
    for i, p in enumerate(C.CONCLUSION):
        body(doc, p, indent=bool(i))

    # --- references -------------------------------------------------------
    heading(doc, "References")
    bib = load_bib()
    missing = [k for k in CITE.order if k not in bib]
    if missing:
        raise SystemExit(f"cited but absent from references.bib: {missing}")
    for i, key in enumerate(CITE.order, 1):
        p = doc.add_paragraph(style="references")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"[{i}]\t{reference_text(bib[key])}")
    return doc


def kpi_table_data():
    """Table III, read from the analysis output rather than retyped."""
    import pandas as pd
    k = pd.read_csv(ROOT / "results" / "kpi_table.csv")
    k = k[k.scenario == "S2"]
    header = ["Indicator", "B1 random", "B2 Industry 4.0",
              "B3 proposed", "\u0394 vs B2", "\u03b4"]
    rows = []
    for _, r in k.iterrows():
        delta = (f"{r.delta_pct:+.1f}%" if r.pct_meaningful
                 else f"{r.delta_abs:+.3f}")
        star = "*" if r.significant else ""
        rows.append([
            r.kpi,
            f"{r.B1_mean:.3f} \u00b1 {r.B1_std:.3f}",
            f"{r.B2_mean:.3f} \u00b1 {r.B2_std:.3f}",
            f"{r.B3_mean:.3f} \u00b1 {r.B3_std:.3f}",
            delta + star,
            f"{r.effect:+.2f}",
        ])
    return header, rows


# =============================================================================
def check_no_first_person(doc) -> None:
    """The article must not speak in the first person. Enforced, not trusted.

    Applied to the prose the author writes, not to the reference list, where a
    cited title may legitimately contain "Our World in Data" and an author's
    middle initial is a capital I. The Roman numeral in "Table I" and the
    lower-case enumerator in "(i)" are likewise excluded rather than being
    reasons to weaken the rule.
    """
    PRONOUNS = re.compile(r"\b(we|us|our|ours|my|mine|ourselves)\b", re.I)
    STANDALONE_I = re.compile(r"(?<!Table )(?<!TABLE )(?<![A-Z]\. )\bI\b(?!\.)")
    PROSE = {"Body Text", "Abstract", "Keywords", "figure caption",
             "table footnote", "Heading 1", "Heading 2"}

    bad = []
    for p in doc.paragraphs:
        if p.style.name not in PROSE:
            continue
        for rx in (PRONOUNS, STANDALONE_I):
            for m in rx.finditer(p.text):
                ctx = p.text[max(0, m.start() - 50):m.end() + 50]
                bad.append(f"{m.group(0)!r}  \u2026{ctx}\u2026")
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for m in PRONOUNS.finditer(c.text):
                    bad.append(f"{m.group(0)!r} in table: {c.text[:60]}")
    if bad:
        print(f"\n  [FAIL] first-person usage found ({len(bad)}):")
        for b in bad[:12]:
            print("        " + b)
        raise SystemExit(1)
    print("  [ok] no first-person pronouns in the prose")


def main() -> None:
    doc = build()
    check_no_first_person(doc)
    doc.save(str(OUT))
    TMP.unlink(missing_ok=True)

    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"  [ok] wrote {OUT.name}")
    print(f"       {len(doc.paragraphs)} paragraphs, ~{words} words, "
          f"{len(doc.tables)} tables, {len(CITE.order)} references")


if __name__ == "__main__":
    main()
