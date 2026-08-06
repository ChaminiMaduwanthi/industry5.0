"""
Build doc.docx — the ten-step plain-Sinhala explanation of the research.

This is the teaching document, not the paper. It explains the work from the
beginning to somebody with no background, using the same analogies and worked
examples that were used when explaining it aloud, and it carries the four
follow-up clarifications that were asked for afterwards.

Sinhala needs a font declared for the complex-script slot as well as the Latin
one, or Word falls back and the text renders as boxes. Every run therefore gets
Nirmala UI on ascii, hAnsi and cs.

Run:  python src/build_summary.py        Writes doc.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "doc.docx"

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

SINHALA = "Nirmala UI"
MONO = "Consolas"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GOOD = RGBColor(0x1E, 0x7A, 0x3C)
WARN = RGBColor(0xB0, 0x3A, 0x2B)


def style_run(run, font=SINHALA, size=11, bold=False, italic=False, color=INK):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(slot), font)
    return run


def para(doc, text="", size=11, bold=False, italic=False, color=INK,
         space_after=6, align=None, font=SINHALA, indent=None):
    p = doc.add_paragraph()
    if text:
        style_run(p.add_run(text), font, size, bold, italic, color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def heading(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(f"{number}   "), SINHALA, 16, True, color=ACCENT)
    style_run(p.add_run(text), SINHALA, 15, True, color=ACCENT)
    bar = doc.add_paragraph()
    bar.paragraph_format.space_after = Pt(8)
    bar.paragraph_format.space_before = Pt(0)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom)
    bar._p.get_or_add_pPr().append(pbdr)
    return p


def sub(doc, text):
    return para(doc, text, size=12, bold=True, color=INK, space_after=4)


def bullet(doc, text, bold_head=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run("•  "), SINHALA, 11, True, color=ACCENT)
    if bold_head:
        style_run(p.add_run(bold_head), SINHALA, 11, True)
    style_run(p.add_run(text), SINHALA, 11)
    return p


def shade(paragraph, hex_fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_fill)
    paragraph._p.get_or_add_pPr().append(sh)


def box(doc, lines, fill="EAF2FA", color=ACCENT, size=11, bold_first=True):
    """A highlighted panel for the ideas that carry the argument."""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(8 if i == 0 else 0)
        p.paragraph_format.space_after = Pt(8 if i == len(lines) - 1 else 2)
        style_run(p.add_run(line), SINHALA, size,
                  bold=(bold_first and i == 0), color=color)
        shade(p, fill)
    return doc


def code(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.space_before = Pt(6 if i == 0 else 0)
        p.paragraph_format.space_after = Pt(6 if i == len(lines) - 1 else 0)
        style_run(p.add_run(line), MONO, 9.5, color=RGBColor(0x22, 0x22, 0x22))
        shade(p, "F4F4F2")


def table(doc, header, rows, widths=None, head_fill="1F4E79"):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = t._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C9D6E4")
        borders.append(el)
    tbl.tblPr.append(borders)

    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        style_run(p.add_run(h), SINHALA, 10, True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(p, head_fill)
        tcpr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), head_fill)
        tcpr.append(sh)

    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            style_run(p.add_run(str(v)), SINHALA, 10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# =============================================================================
#  the ten steps — Sinhala
# =============================================================================
def build_si() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)

    # ---- title ----------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    style_run(t.add_run("මගේ පර්යේෂණය — මුල ඉඳන්"), SINHALA, 22, True, color=ACCENT)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(2)
    style_run(st.add_run("Industry 5.0 සඳහා මිනිසා සහ යන්ත්‍රය එකට ආකෘතිගත කරන "
                         "Digital Twin රාමුවක්"), SINHALA, 12, color=MUTED)

    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.paragraph_format.space_after = Pt(14)
    style_run(a.add_run("Chamini Maduwanthi  ·  පියවර 10කින්, සරලව"),
              SINHALA, 10.5, italic=True, color=MUTED)

    box(doc, ["මේක කාටද?",
              "පරිගණක විද්‍යාව හෝ කර්මාන්ත ඉංජිනේරු විද්‍යාව ගැන කිසිවක් නොදන්නා "
              "කෙනෙකුට තේරෙන විදිහට ලියා ඇත. පියවර 10ම පිළිවෙලට කියවන්න."],
        fill="EAF2FA")

    # ---- 1 ---------------------------------------------------------------
    heading(doc, "1️⃣", "ප්‍රශ්නය මොකක්ද?")
    para(doc, "කර්මාන්තශාලාවක් හිතන්න. යන්ත්‍ර 5ක්. සේවකයෝ 3ක්.")
    para(doc, "හැම විනාඩි 15කට වරක්ම කවුරුහරි තීරණයක් ගන්න ඕන:")
    box(doc, ["“මේ අලුත් වැඩේ කාට දෙනවද? කුමන යන්ත්‍රයටද?”"],
        fill="F4F4F2", color=INK)
    para(doc, "සරලයි වගේ පේනවා. ඒත් ඒ එක තීරණයෙන් එකවර දේවල් තුනක් තීරණය වෙනවා:")
    code(doc, ["👷   ඒ මිනිසා කොච්චර වෙහෙසෙනවද",
               "🌱   කොච්චර විදුලියක් යනවද, කොච්චර නාස්ති වෙනවද",
               "💰   කොච්චර නිෂ්පාදනය වෙනවද"])
    box(doc, ["දැන් තියෙන පරිගණක ක්‍රම බලන්නේ අන්තිම එක විතරයි — නිෂ්පාදනය.",
              "මිනිසා ගැන බලන්නේ නෑ."], fill="FDEEEA", color=WARN)

    sub(doc, "ඇයි විනාඩි 15?")
    para(doc, "බත් හැළිය කොච්චර වෙලාවකට වරක් බලනවද? තත්පර 10කට වරක් නම් නාස්තියි — "
              "කිසිවක් වෙනස් වෙලා නෑ. පැයකට වරක් නම් පිච්චිලා ඉවරයි. "
              "විනාඩි 10–15 හරි.")
    para(doc, "අපේ කර්මාන්තශාලාවේත් එහෙමයි: වැඩක් ගන්න විනාඩි 8–18ක් යනවා, සහ "
              "වෙහෙස වෙනස් වෙන්නත් විනාඩි 15ක් විතර යනවා. විනාඩි 1ක් නම් නාස්තියි; "
              "පැයක් නම් මිනිහෙක් වෙහෙසිලා ඉවර වුණත් අපි දැනගන්නේ පරක්කු වෙලා.",
         space_after=4)
    para(doc, "පැය 8 = විනාඩි 480.  480 ÷ 15 = දවසකට තීරණ 32ක්.",
         italic=True, color=MUTED)

    # ---- 2 ---------------------------------------------------------------
    heading(doc, "2️⃣", "Industry 4.0 සහ Industry 5.0")
    table(doc, ["", "අදහස"],
          [["Industry 4.0", "“හැම දෙයක්ම automatic කරමු, ඉක්මන් කරමු” "
                            "— කාර්යක්ෂමතාවය"],
           ["Industry 5.0", "“හොඳයි… ඒත් කාටද ඒක? මිනිසාට හොඳද? "
                            "පරිසරයට හොඳද?”"]],
          widths=[1.3, 5.0])
    para(doc, "යුරෝපා කොමිසමේ නිල නිර්වචනය අනුව Industry 5.0 හි කුලුනු තුනක් තියෙනවා: "
              "මිනිසා · පරිසරය · ඔරොත්තු දීම.")
    box(doc, ["උපමාව",
              "Industry 4.0 = වේගවත්ම කාර් එක.",
              "Industry 5.0 = වේගවත්, ඒත් රියදුරාට ආරක්ෂිත, ඉන්ධන අඩු කාර් එක."],
        fill="EAF2FA")

    # ---- 3 ---------------------------------------------------------------
    heading(doc, "3️⃣", "Digital Twin කියන්නේ මොකක්ද?")
    box(doc, ["යමකගේ සජීවී පරිගණක පිටපතක්."], fill="EAF2FA")
    code(doc, ["ඇත්ත යන්ත්‍රය   ←→   පරිගණකයේ ඒකේ පිටපත",
               "                     (sensor වලින් හැම විටම update වෙනවා)"])
    para(doc, "උදාහරණය: ඔබේ ෆෝන් එකේ battery indicator එක. ඒක බැටරියේ digital twin "
              "එකක් — ඇත්ත බැටරියේ තත්ත්වය තිරයේ පෙන්නනවා. පිටපත බලලා "
              "“charge කරන්න ඕන” කියලා තීරණයක් ගන්න පුළුවන්.")

    # ---- 4 ---------------------------------------------------------------
    heading(doc, "4️⃣", "අපේ පරතරය — කවුරුවත් නොකළ දේ")
    para(doc, "Digital twin ගැන පර්යේෂණ ගොඩක් තියෙනවා. අපි studies 46ක් කියෙව්වා. "
              "ක්ෂේත්‍රය දෙකට බෙදිලා තියෙනවා:")
    code(doc, ["කණ්ඩායම A  ─  යන්ත්‍රයට twin එකක් හදනවා                    ✅",
               "              මිනිසාට “නම, වයස, දක්ෂතාව” වගේ ස්ථිතික දේ    ❌",
               "",
               "කණ්ඩායම B  ─  මිනිසාගේ වෙහෙස මනිනවා                        ✅",
               "              ඒත් ඒක බලලා තීරණයක් ගන්නේ නෑ                  ❌"])
    box(doc, ["studies 46න් එකක්වත් තුනම එකට කරලා නෑ:",
              "①  මිනිසාගේ සජීවී තත්ත්වය    ②  යන්ත්‍රයේ twin එක    "
              "③  මනින ලද පාරිසරික ඉලක්ක",
              "ඒක තමයි අපේ පරතරය."], fill="EAF2FA")

    # ---- 5 ---------------------------------------------------------------
    heading(doc, "5️⃣", "අපි හදපු දේ — කෑලි 4කින්")

    sub(doc, "🔧  කෑල්ල 1 — යන්ත්‍රයේ Twin එක")
    para(doc, "යන්ත්‍රයක් ගැන දේවල් 5ක් මතක තියාගන්නවා: සෞඛ්‍යය · විදුලිය · "
              "දෝෂ අවදානම · ලබාගත හැකිද · භාවිතය.")
    code(doc, ["යන්ත්‍රයක් ඉවර වෙන්නේ 216 විනාඩියක වැඩෙන් (ඇත්ත දත්තවලින්)",
               "",
               "සැහැල්ලු වැඩ දුන්නොත්  →  291 විනාඩි යනවා     🙂",
               "බර වැඩ දුන්නොත්        →  145 විනාඩි විතරයි   😰   දෙගුණයක් ඉක්මනට"])

    sub(doc, "👷  කෑල්ල 2 — මිනිසාගේ Twin එක   ★ මේකයි අලුත් දේ")
    para(doc, "අනිත් අය වෙහෙස ස්ථිර සංඛ්‍යාවක් විදිහට සලකනවා: "
              "“බර වැඩේට වෙහෙස = 8” — ඉවරයි.")
    box(doc, ["ඒත් ඒක වැරදියි:",
              "උදේ 8ට බර වැඩක් කරන කෙනා   ≠   හවස 4ට එකම බර වැඩ කරන කෙනා"],
        fill="FDEEEA", color=WARN)
    para(doc, "අපි වෙහෙස ජලය පිරෙන බාල්දියක් වගේ සලකනවා: වැඩ කරනකොට පිරෙනවා, "
              "විවේක ගන්නකොට හිස් වෙනවා.")
    para(doc, "ඊටත් වඩා වැදගත්: හැමෝගේම බාල්දිය එකම ප්‍රමාණය නෙවෙයි.", bold=True)
    para(doc, "වයස, බර, උස, ස්ත්‍රී/පුරුෂ අනුව “දිගටම දරාගන්න පුළුවන් උපරිමය” "
              "ගණනය කරනවා — ප්‍රකාශිත වෛද්‍ය සමීකරණවලින් (Mifflin 1990, Price 1990).")
    table(doc, ["සේවකයා", "වයස", "දරාගත හැකි උපරිමය",
                "මධ්‍යම වැඩ (4.26) දිගටම කරන්න පුළුවන්ද?"],
          [["OP1  පිරිමි", "28", "5.35", "✅  පුළුවන්"],
           ["OP3  පිරිමි", "47", "4.55", "✅  පුළුවන් (ලං වෙලා)"],
           ["OP2  ගැහැනු", "35", "3.62", "⛔  බෑ!"]],
          widths=[1.2, 0.6, 1.5, 3.0])
    box(doc, ["මේක තමයි ලස්සනම කොටස.",
              "OP2 ට medium වැඩ දිගටම කරන්න බැරි බව අපි කොහෙවත් ලියලා නෑ. "
              "ඒක ඇගේ ශරීරයේ සංඛ්‍යාවලින්ම එනවා. System එක තනිවම එයාට වැඩිපුර "
              "විවේක දෙනවා.",
              "එකම නීතියක් — ඒත් හැමෝටම වෙනස් ආරක්ෂාවක්."], fill="E9F5EC",
        color=GOOD)

    para(doc, "“4.26” කියන්නේ මොකක්ද? විනාඩියකට ශරීරය පුළුස්සන කැලරි ගණන. "
              "ෆෝන් එකේ බැටරිය බහින වේගය වගේ. ඒක ISO 8996 කියන ජාත්‍යන්තර "
              "ප්‍රමිතියෙන් එනවා — අපි හදපු අංකයක් නෙවෙයි.", space_after=4)
    table(doc, ["වැඩේ", "ISO පන්තිය", "kcal/min", "පැය 8ක shift එකකට"],
          [["සැහැල්ලු", "Class 1", "2.58", "1,239 kcal"],
           ["මධ්‍යම", "Class 2", "4.26", "2,044 kcal"],
           ["බර", "Class 3", "5.94", "2,850 kcal"]],
          widths=[1.3, 1.2, 1.2, 1.8])

    sub(doc, "🔗  කෑල්ල 3 — Twin දෙක එකිනෙකට කතා කරනවා")
    code(doc, ["මිනිසා → යන්ත්‍රය :  වෙහෙසුණු කෙනා     →  වැඩිපුර දෝෂ",
               "                     දක්ෂතාව අඩු කෙනා   →  වැඩිපුර දෝෂ",
               "",
               "යන්ත්‍රය → මිනිසා :  ගෙවුණු යන්ත්‍රයක්   →  වැඩි මානසික බර",
               "                     බර වැඩක්           →  වේගෙන් වෙහෙසෙනවා",
               "                     වේගවත් යන්ත්‍රයක්   →  නරක ඉරියව්ව"])
    para(doc, "උදාහරණය: වෙහෙසුණු කෙනෙක් වැරදි කරනවා → නිෂ්පාදනය නරක් වෙනවා → "
              "ඒක යන්ත්‍රයේ ප්‍රතිඵලයක් වගේ පේනවා. ඇත්තටම ඒක මිනිසාගේ ප්‍රතිඵලයක්. "
              "අපේ system එකට ඒක පේනවා.")

    sub(doc, "⚖️  කෑල්ල 4 — තීරණය ගන්න විදිහ   ★★★ ප්‍රධානම අදහස")
    para(doc, "ක්‍රම දෙකක් තියෙනවා.", bold=True)
    para(doc, "❌  ක්‍රමය 1 — “දඩය” (soft penalty).  අනිත් අය කරන්නේ මේක:",
         space_after=2)
    code(doc, ["නිෂ්පාදනය වැඩියි    →  +100 ලකුණු",
               "මිනිසා වෙහෙසෙනවා   →   −30 ලකුණු  (දඩයක්)",
               "                      ─────────",
               "එකතුව              =  +70   →  “හොඳයි, කරමු!”"])
    box(doc, ["ප්‍රශ්නය: නිෂ්පාදනයේ ලාභය ලොකු නම්, වෙහෙසේ දඩය හැමවිටම ගෙවන්න "
              "පුළුවන්. මිනිසා විකුණන්න පුළුවන් දෙයක් වෙනවා."],
        fill="FDEEEA", color=WARN)
    para(doc, "✅  ක්‍රමය 2 — “පෙරහන” (hard constraint).  අපි කරන්නේ මේක:",
         space_after=2)
    code(doc, ["පියවර 1 :  මිනිසාට හානියක් වෙන විකල්ප ඔක්කොම අයින් කරන්න   ← මුලින්ම",
               "පියවර 2 :  ඉතුරු ඒවා අතරින් හොඳම එක තෝරන්න"])
    box(doc, ["වෙනස",
              "පෙරහනකින් අයින් වුණු විකල්පයක් කිසිම මිලකට ආපහු ගන්න බෑ. "
              "ලාභය කොච්චර ලොකු වුණත් නෑ.",
              "මත්පැන් බීලා රිය පදවන එකට දඩයක් නෙවෙයි තියෙන්නේ — ඒක තහනම්. "
              "එහෙම තමයි අපි වෙහෙසට සලකන්නේ."], fill="E9F5EC", color=GOOD)
    table(doc, ["", "නීතිය", "තේරුම", "අගය කොහෙන්ද?"],
          [["HC1", "වෙහෙස < 80%", "තමන්ගේ සීමාවෙන් 80% පැනීම තහනම්", "අපේ *"],
           ["HC2", "දක්ෂතාව ≥ 0.40", "නොදන්නා වැඩක් දෙන එක තහනම්", "අපේ"],
           ["HC3", "ඉරියව් අවදානම ≤ 5", "භයානක ඉරියව්වක් තහනම්",
            "✅ සාහිත්‍යයෙන් (RULA)"],
           ["HC4", "යන්ත්‍ර සෞඛ්‍යය > 0.30", "කැඩෙන්න ළං යන්ත්‍රයක් තහනම්", "අපේ"]],
          widths=[0.5, 1.5, 2.6, 1.6])
    para(doc, "* “80%” කියන අංකය අපේ. ඒත් 80% කුමකින්ද කියන එක "
              "සාහිත්‍යයෙන් — ඒක “දිගටම දරාගන්න පුළුවන් උපරිමයෙන්” 80%. "
              "වේග සීමාව 60 km/h කියලා දාන එක තෝරාගැනීමක්; ඒත් “km/h” "
              "කියන මිනුම තෝරාගැනීමක් නෙවෙයි.", size=10, italic=True, color=MUTED)
    para(doc, "සහ අපි ඒවා පරීක්ෂා කළා: HC1 එක 0.70 · 0.80 · 0.90 කරලා බැලුවා — "
              "තුනේම එකම නිගමනය. HC4 එක අයින් කළත් කිසිම වෙනසක් නෑ. "
              "HC1 සහ HC3 අයින් කළොත් ඇත්තටම මිනිසුන්ට හානි වෙනවා.", size=10,
         italic=True, color=MUTED)

    # ---- 6 ---------------------------------------------------------------
    heading(doc, "6️⃣", "පරීක්ෂා කරපු විදිහ")
    para(doc, "ප්‍රශ්නය: ලෝකයේ යන්ත්‍ර දත්තත් මිනිස් දත්තත් එකට තියෙන කර්මාන්තශාලා "
              "dataset එකක් නෑ.")
    para(doc, "විසඳුම: simulation එකක් (පරිගණකයේ කර්මාන්තශාලාවක්) — "
              "ඒත් සංඛ්‍යා ඇත්ත දත්තවලින්:")
    code(doc, ["යන්ත්‍රයේ ආයුෂ     ←  ඇත්ත maintenance dataset එකකින්",
               "විදුලි පරිභෝජනය    ←  ඇත්ත වානේ කර්මාන්තශාලාවක දත්තවලින්",
               "දෝෂ අනුපාතය       ←  ඇත්ත semiconductor නිෂ්පාදන දත්තවලින්",
               "මිනිසාගේ ශරීරය     ←  ප්‍රකාශිත වෛද්‍ය සමීකරණවලින්"])
    table(doc, ["ක්‍රමය", "විස්තරය"],
          [["B1", "අහඹු ලෙස වැඩ බෙදනවා (පහළම මට්ටම)"],
           ["B2", "Industry 4.0 — යන්ත්‍ර twin එක පූර්ණ බලයෙන්, "
                  "ඒත් මිනිසා ගැන බලන්නේ නෑ"],
           ["B3", "අපේ එක — twin දෙකම + තහනම් නීති"]],
          widths=[0.8, 5.5])
    code(doc, ["තත්ත්ව 3ක් :  සාමාන්‍ය  ·  ඉල්ලුම 150%  ·  යන්ත්‍ර කැඩීම",
               "ධාවනය      :  3 × 3 × seeds 30  =  runs 270"])
    box(doc, ["B2 එක දුර්වල කරන්න බෑ.",
              "ඒක “පිදුරු මිනිසෙක්” නම් ඒක වංචාවක්. ඒ නිසා B2 එකට "
              "යන්ත්‍ර twin එකේ සම්පූර්ණ බලය දුන්නා. එකම වෙනස: මිනිසා ගැන "
              "හිතන්නේ නෑ."], fill="EAF2FA")

    # ---- 7 ---------------------------------------------------------------
    heading(doc, "7️⃣", "ප්‍රතිඵල — එකින් එක")
    para(doc, "ඉල්ලුම වැඩි තත්ත්වයේ, අපේ එක vs Industry 4.0:", bold=True)
    table(doc, ["ප්‍රතිඵලය", "Industry 4.0", "අපේ එක", "වෙනස"],
          [["වෙහෙස", "0.741", "0.541", "▼ 27%"],
           ["නීති කැඩීම (shift එකකට)", "79.5", "0.0", "▼ 100%"],
           ["විදුලිය ඒකකයකට", "1.321", "0.915", "▼ 31%"],
           ["CO₂", "39.9 kg", "27.2 kg", "▼ 32%"],
           ["යන්ත්‍ර නැවතුම්", "1.858", "1.358", "▼ 27%"],
           ["නිෂ්පාදනය", "91.9", "90.2", "▼ 1.8%  ← වැදගත්"]],
          widths=[2.2, 1.4, 1.2, 1.6])
    box(doc, ["★★★  පත්‍රිකාවේ ප්‍රධානම වාක්‍යය",
              "ඉල්ලුම වැඩිම වෙලාවේ — Industry 4.0 ක්‍රමය මිනිසාව වැඩිපුරම හිර "
              "කරන වෙලාවේ — මිනිසා රැක ගැනීමට මනින්න පුළුවන් නිෂ්පාදන අලාභයක් "
              "වුණේ නෑ.",
              "ඒ 1.8% අඩුවීමේ p = 0.148. ඒ කියන්නේ ඒක අහම්බයකින් වෙන්න පුළුවන් "
              "— ඇත්ත වෙනසක් නොවෙන්න පුළුවන්."], fill="E9F5EC", color=GOOD)
    para(doc, "සහ තවත් එකක්: යන්ත්‍රත් අඩුවෙන් නවතිනවා (−27%). ඒ කියන්නේ "
              "“මිනිසාට හොඳ දේ යන්ත්‍රයටත් හොඳයි”. වෙහෙසුණු කෙනා වැරදි "
              "කරනවා → යන්ත්‍රයට බර වැඩියි.")

    # ---- 8 ---------------------------------------------------------------
    heading(doc, "8️⃣", "බලාපොරොත්තු නොවූ ලොකුම සොයාගැනීම")
    para(doc, "අපි “ප්‍රමුඛතා” 4ක් හැදුවා — මිනිසාට, පරිසරයට, ලාභයට, "
              "සමානව. ඒවා හතර හාත්පසින්ම වෙනස්.")
    code(doc, ["W-Balanced :  වෙහෙස 0.541      නිෂ්පාදනය 90.2",
               "W-Human    :  වෙහෙස 0.541      නිෂ්පාදනය 90.3",
               "W-Green    :  වෙහෙස 0.542      නිෂ්පාදනය 90.3",
               "W-Profit   :  වෙහෙස 0.542      නිෂ්පාදනය 90.3",
               "                 ↑                    ↑",
               "             හතරම එකයි!          හතරම එකයි!"])
    para(doc, "ඇයි? ගණන් කරලා බැලුවා:", bold=True)
    code(doc, ["තහනම් නීති පෙරලාට පස්සේ...",
               "තීරණ 45,754න්  →  95.7%ට  ඉතුරු වෙන්නේ  විකල්ප 0 හෝ 1යි!"])
    box(doc, ["උපමාව — කෑම කඩේ",
              "ඔබට කුකුල් මස් ආසයි, මාළු එපා. ඒත් කඩේ නීති යෙදුවම menu එකේ "
              "ඉතුරු වෙන්නේ එක් කෑමක් විතරයි.",
              "දැන් ඔබේ කැමැත්තෙන් වැඩක් තියෙනවද? නෑ. තියෙන්නේ එකයි.",
              "තෝරන්න දෙයක් නෑ නම්, කැමැත්තෙන් වැඩක් නෑ."], fill="EAF2FA")
    box(doc, ["මේක හොඳ ආරංචියක් — නරක එකක් නෙවෙයි",
              "ප්‍රමුඛතාවක් වෙනස් කරලා මිනිසාගේ ආරක්ෂාව අයින් කරන්න බෑ. "
              "ලාභයට 100% ප්‍රමුඛතාවය දුන්නත් වෙහෙසේ තහනම තාමත් තියෙනවා.",
              "“මේ framework එක මෙහෙයවන්නේ තහනම් නීතියි — ප්‍රමුඛතා නෙවෙයි.”"],
        fill="E9F5EC", color=GOOD)

    # ---- 9 ---------------------------------------------------------------
    heading(doc, "9️⃣", "අවංකව කිව්ව දුර්වලතා")
    para(doc, "මේකයි හොඳ පර්යේෂණයක ලකුණ — අඩුපාඩු හංගන්නේ නෑ.", italic=True,
         color=MUTED)
    table(doc, ["#", "දුර්වලතාව", "ඇයි ලිව්වේ"],
          [["1", "ඉරියව් අවදානම 20%කින් අඩු වුණේ නෑ (13–15% විතරයි)",
            "අපේම උපකල්පනය සම්පූර්ණ වුණේ නෑ — ඒත් ඒකට ගණිතමය හේතුවක් තියෙනවා"],
           ["2", "වැඩ බෙදීම අසමානයි", "අපිට විරුද්ධ ප්‍රතිඵලයක් — ඒත් හංගන්නේ නෑ"],
           ["3", "වේග දෙකක් (λ, μ) මනිලා නෑ, කැලිබ්‍රේට් කරලා",
            "0.5×–2× කරලා නිගමනය වෙනස් වෙන්නේ නෑ කියලා පෙන්නලා"],
           ["4", "පැහැදිලි කිරීමේ ස්ථරය (L5) ලියලා නෑ",
            "Architecture එකේ තියෙනවා, ලියලා නෑ — ඒක කියලා තියෙනවා"],
           ["5", "ඇත්ත කර්මාන්තශාලාවක පරීක්ෂා කරලා නෑ",
            "Simulation එකක් — ඒක පැහැදිලිව කියනවා"]],
          widths=[0.35, 2.6, 3.4])
    box(doc, ["#2 එකේ ලස්සන දෙයක් තියෙනවා",
              "වැඩ අසමානව බෙදෙන්නේ ඇයි කියලා හෙව්වම හම්බුණේ: OP1 (ශක්තිමත්ම කෙනා) "
              "ට බර වැඩ තහනම් — ඔහුගේ දක්ෂතාව 0.30, සීමාව 0.40. ඒ නිසා බර වැඩ "
              "යන්නේ දරාගන්න අඩුම හැකියාව ඇති දෙන්නාට.",
              "ඒ එක දක්ෂතාව පුහුණු කළොත් පරතරයෙන් 2/3ක් වැහෙනවා, නිෂ්පාදනයත් "
              "වැඩි වෙනවා.",
              "යන්ත්‍රය විතරක් බලන system එකකට මේ ප්‍රශ්නය අහන්නවත් බෑ — "
              "ඒකේ මිනිසාගේ ආකෘතියක්ම නෑ."], fill="EAF2FA")

    # ---- 10 --------------------------------------------------------------
    heading(doc, "🔟", "මුළු දේම — වාක්‍ය 3කින්")
    for n, line in enumerate([
            "Industry 5.0 මිනිසා මධ්‍යයට ගෙන එනවා. ඒත් digital twins තවම "
            "යන්ත්‍ර ගැන විතරයි. මිනිසා ස්ථිතික දත්තයක් විතරයි.",
            "අපි මිනිසාගේ වෙහෙස සජීවී තත්ත්වයක් කරලා, යන්ත්‍රයේ twin එකට "
            "සම්බන්ධ කරලා, මිනිසාගේ සීමා “දඩයක්” නොව “තහනමක්” කළා.",
            "ප්‍රතිඵලය: වෙහෙස −27%, විදුලිය −31%, නීති කැඩීම 79.5 → 0 — "
            "මනින්න පුළුවන් නිෂ්පාදන අලාභයක් නැතුව."], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(8)
        style_run(p.add_run(f"{n}.  "), SINHALA, 13, True, color=ACCENT)
        style_run(p.add_run(line), SINHALA, 12)

    doc.add_paragraph()
    box(doc, ["තව දැනගන්න:",
              "paper.docx — සම්පූර්ණ පත්‍රිකාව  ·  "
              "docs/14-demo-guide.md — supervisor ට පෙන්නන හැටි  ·  "
              "docs/12-paper-blueprint.md — හැම විස්තරයක්ම"],
        fill="F4F4F2", color=MUTED, size=10)
    return doc


# =============================================================================
#  the ten steps — English
# =============================================================================
EN = "Segoe UI"


def build_en() -> Document:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)

    def P(text="", **kw):
        kw.setdefault("font", EN)
        return para(doc, text, **kw)

    def H(number, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        style_run(p.add_run(f"{number}   "), EN, 16, True, color=ACCENT)
        style_run(p.add_run(text), EN, 15, True, color=ACCENT)
        bar = doc.add_paragraph()
        bar.paragraph_format.space_after = Pt(8)
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1F4E79")
        pbdr.append(bottom)
        bar._p.get_or_add_pPr().append(pbdr)

    def S(text):
        return para(doc, text, size=12, bold=True, space_after=4, font=EN)

    def B(lines, fill="EAF2FA", color=ACCENT, size=11):
        for i, line in enumerate(lines):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(8 if i == 0 else 0)
            p.paragraph_format.space_after = Pt(8 if i == len(lines) - 1 else 2)
            style_run(p.add_run(line), EN, size, bold=(i == 0), color=color)
            shade(p, fill)

    def T(header, rows, widths=None):
        t = table(doc, header, rows, widths)
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.font.color and r.font.color.rgb == RGBColor(
                                0xFF, 0xFF, 0xFF):
                            style_run(r, EN, 10, True,
                                      color=RGBColor(0xFF, 0xFF, 0xFF))
                        else:
                            style_run(r, EN, 10)
        return t

    # ---- title ----------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    style_run(t.add_run("My Research — From the Beginning"), EN, 22, True,
              color=ACCENT)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(2)
    style_run(st.add_run("A digital twin framework that models the machine and "
                         "the person together, for Industry 5.0"), EN, 12,
              color=MUTED)
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.paragraph_format.space_after = Pt(14)
    style_run(a.add_run("Chamini Maduwanthi   ·   explained in ten steps"),
              EN, 10.5, italic=True, color=MUTED)

    B(["Who is this for?",
       "Anyone with no background in computing or industrial engineering. "
       "Read the ten steps in order; each one builds on the last."])

    # ---- 1 ---------------------------------------------------------------
    H("1", "What is the problem?")
    P("Picture a small factory. Five machines. Three workers.")
    P("Every fifteen minutes, somebody has to make a decision:")
    B(["“This new job — who should do it, and on which machine?”"],
      fill="F4F4F2", color=INK)
    P("It sounds simple. But that single decision settles three things at once:")
    code(doc, ["PEOPLE   how tired that person becomes",
               "PLANET   how much electricity is used, how much is wasted",
               "PROFIT   how much gets produced"])
    B(["Today's scheduling systems look at the third one only — output.",
       "They do not look at the person at all."], fill="FDEEEA", color=WARN)

    S("Why fifteen minutes?")
    P("Think of checking a pot of rice. Every ten seconds is wasted effort — "
      "nothing has changed. Once an hour and it has burnt. Every ten or "
      "fifteen minutes is about right.")
    P("The factory is the same. A job takes eight to eighteen minutes, and "
      "tiredness takes about fifteen minutes to change noticeably. Deciding "
      "every minute would be pointless; deciding once an hour would mean "
      "somebody passes their limit and nobody notices until far too late.",
      space_after=4)
    P("An eight-hour shift is 480 minutes. 480 ÷ 15 = 32 decisions per day.",
      italic=True, color=MUTED)

    # ---- 2 ---------------------------------------------------------------
    H("2", "Industry 4.0 and Industry 5.0")
    T(["", "The idea"],
      [["Industry 4.0", "“Automate everything, make it faster.” — efficiency"],
       ["Industry 5.0", "“Fine… but who is it for? Is it good for the person? "
                        "For the planet?”"]],
      widths=[1.3, 5.0])
    P("The European Commission's official definition gives Industry 5.0 three "
      "pillars: people, sustainability and resilience.")
    B(["An analogy",
       "Industry 4.0 = the fastest car.",
       "Industry 5.0 = a fast car that is also safe for the driver and uses "
       "less fuel."])

    # ---- 3 ---------------------------------------------------------------
    H("3", "What is a digital twin?")
    B(["A live computer copy of something real."])
    code(doc, ["the real machine   <-->   its copy inside the computer",
               "                          (updated constantly from sensors)"])
    P("Example: the battery indicator on your phone. That is a digital twin of "
      "the battery — the real battery's state, shown on screen. You look at "
      "the copy and decide “time to charge”. You never open the phone.")

    # ---- 4 ---------------------------------------------------------------
    H("4", "The gap — what nobody had done")
    P("There is plenty of research on digital twins. Reading 46 studies showed "
      "the field is split in two:")
    code(doc, ["Group A  -  builds a twin of the machine                    YES",
               "            represents the person by fixed facts:",
               "            name, age, skill level                          NO",
               "",
               "Group B  -  measures the person's tiredness                 YES",
               "            but never uses it to make a decision            NO"])
    B(["Not one of the 46 studies does all three at once:",
       "(1) the person's live state    (2) a machine twin    "
       "(3) measured environmental goals",
       "That is the gap this research fills."])

    # ---- 5 ---------------------------------------------------------------
    H("5", "What was built — in four pieces")

    S("Piece 1 — the machine's twin")
    P("Each machine keeps track of five things: health, energy use, defect "
      "risk, whether it is available, and how busy it has been.")
    code(doc, ["a machine wears out after 216 minutes of work",
               "                          (taken from real maintenance data)",
               "",
               "give it light work  ->  291 minutes           :)",
               "give it heavy work  ->  145 minutes           :(   twice as fast"])

    S("Piece 2 — the person's twin        <-- this is the new part")
    P("Other researchers treat tiredness as a fixed number: “heavy job = 8 "
      "points of strain”. That is the end of it.")
    B(["But that is wrong:",
       "a person doing heavy work at 8 in the morning is not the same as the "
       "same person doing the same work at 4 in the afternoon"],
      fill="FDEEEA", color=WARN)
    P("Here, tiredness is treated like a bucket filling with water. Working "
      "fills it. Resting empties it.")
    P("And more important still: not everybody's bucket is the same size.",
      bold=True)
    P("From age, weight, height and sex, the model computes the work rate each "
      "person can sustain all day — using published medical equations "
      "(Mifflin 1990, Price 1990).")
    T(["Worker", "Age", "Sustainable maximum",
       "Can they do medium work (4.26) all day?"],
      [["OP1  male", "28", "5.35", "Yes"],
       ["OP3  male", "47", "4.55", "Yes, but close to the line"],
       ["OP2  female", "35", "3.62", "No"]],
      widths=[1.2, 0.6, 1.6, 2.9])
    B(["This is the best part.",
       "Nowhere does the system say “OP2 cannot do medium work”. That falls "
       "out of the numbers describing her body. The system gives her extra "
       "rest on its own.",
       "One rule for everyone — but different protection for each person."],
      fill="E9F5EC", color=GOOD)
    P("What is “4.26”? It is calories burnt per minute — the rate the body is "
      "being drained, like the rate a phone battery discharges. It comes from "
      "ISO 8996, an international standard, not from us.", space_after=4)
    T(["Work", "ISO class", "kcal/min", "Over an 8-hour shift"],
      [["Light", "Class 1", "2.58", "1,239 kcal"],
       ["Medium", "Class 2", "4.26", "2,044 kcal"],
       ["Heavy", "Class 3", "5.94", "2,850 kcal"]],
      widths=[1.3, 1.2, 1.2, 1.8])

    S("Piece 3 — the two twins talk to each other")
    code(doc, ["person  -> machine :  a tired worker      ->  more defects",
               "                      a less skilled one  ->  more defects",
               "",
               "machine -> person  :  a worn machine      ->  more mental load",
               "                      a heavy job         ->  tires faster",
               "                      a fast machine      ->  worse posture"])
    P("Example: a tired worker makes mistakes, so the product comes out "
      "faulty. That looks like a machine problem. It is really a people "
      "problem. This system can see that.")

    S("Piece 4 — how the decision is made        <-- the central idea")
    P("There are two ways to do it.", bold=True)
    P("Way 1 — a “fine” (a penalty). This is what others do:", space_after=2)
    code(doc, ["high output          ->  +100 points",
               "the worker gets tired ->   -30 points   (the fine)",
               "                         ---------",
               "total                 =   +70   ->  “good enough, do it”"])
    B(["The problem: if the profit is large enough, the fine on tiredness can "
       "always be paid. The person becomes something you can trade away."],
      fill="FDEEEA", color=WARN)
    P("Way 2 — a “filter” (a hard constraint). This is what was built:",
      space_after=2)
    code(doc, ["step 1 :  remove every option that would harm the person   <- first",
               "step 2 :  pick the best of whatever is left"])
    B(["The difference",
       "An option removed by the filter cannot be bought back at any price, "
       "no matter how large the profit.",
       "Drink-driving does not carry a fine you can pay to proceed — it is "
       "forbidden. Tiredness is treated the same way here."],
      fill="E9F5EC", color=GOOD)
    T(["", "The rule", "In plain words", "Where the number comes from"],
      [["HC1", "tiredness < 80%", "cannot go past 80% of your own limit",
        "ours *"],
       ["HC2", "skill >= 0.40", "cannot give someone a job they cannot do",
        "ours"],
       ["HC3", "posture risk <= 5", "cannot use a dangerous posture",
        "from the literature (RULA)"],
       ["HC4", "machine health > 0.30", "cannot use a machine about to fail",
        "ours"]],
      widths=[0.5, 1.4, 2.5, 1.8])
    P("* The number 80% is ours. But what it is 80% of comes from the "
      "literature — it is 80% of the rate that person can sustain all day. "
      "Setting a speed limit at 60 km/h is a choice; the unit “km/h” is not.",
      size=10, italic=True, color=MUTED)
    P("And they were tested. HC1 was rerun at 0.70, 0.80 and 0.90 — the same "
      "conclusion every time. Removing HC4 changed nothing at all. Removing "
      "HC1 or HC3 caused real harm to the workers.",
      size=10, italic=True, color=MUTED)

    # ---- 6 ---------------------------------------------------------------
    H("6", "How it was tested")
    P("The obstacle: no public dataset in the world contains machine data and "
      "human data from the same factory.")
    P("The answer: a simulation — a factory inside the computer — but with "
      "every number taken from real data:")
    code(doc, ["machine lifetime      <-  a real maintenance dataset",
               "electricity use       <-  a real steel plant's records",
               "defect rate           <-  a real semiconductor production line",
               "the human body        <-  published medical equations"])
    T(["Method", "What it does"],
      [["B1", "assigns work at random (a floor, for sanity)"],
       ["B2", "Industry 4.0 — full machine twin, but blind to the person"],
       ["B3", "this research — both twins, plus the forbidden rules"]],
      widths=[0.8, 5.5])
    code(doc, ["three situations :  normal  ·  demand at 150%  ·  machines break",
               "runs             :  3 x 3 x 30 random seeds  =  270 runs"])
    B(["B2 was not allowed to be weak.",
       "A straw opponent would make the whole comparison worthless. B2 keeps "
       "the machine twin at full strength. The only difference is that it does "
       "not think about the person."])

    # ---- 7 ---------------------------------------------------------------
    H("7", "The results, one by one")
    P("Under high demand, this research against Industry 4.0:", bold=True)
    T(["Measure", "Industry 4.0", "This work", "Change"],
      [["Tiredness", "0.741", "0.541", "down 27%"],
       ["Rules broken (per shift)", "79.5", "0.0", "down 100%"],
       ["Electricity per unit", "1.321", "0.915", "down 31%"],
       ["Carbon", "39.9 kg", "27.2 kg", "down 32%"],
       ["Machine stoppages", "1.858", "1.358", "down 27%"],
       ["Output", "91.9", "90.2", "down 1.8%  <- the key one"]],
      widths=[2.2, 1.4, 1.2, 1.6])
    B(["The single most important sentence in the paper",
       "At the moment of highest demand — exactly when a throughput-driven "
       "scheduler pushes people hardest — protecting the workers cost no "
       "measurable production.",
       "That 1.8% has p = 0.148, which means it could easily be chance rather "
       "than a real difference."], fill="E9F5EC", color=GOOD)
    P("And one more: the machines stopped less often too (down 27%). What is "
      "good for the person turns out to be good for the equipment — a tired "
      "worker makes mistakes, and the machine absorbs the consequences.")

    # ---- 8 ---------------------------------------------------------------
    H("8", "The most surprising finding")
    P("Four different sets of priorities were tried — favouring the person, "
      "the planet, profit, and a balance. They differ enormously.")
    code(doc, ["balanced     :  tiredness 0.541      output 90.2",
               "people first :  tiredness 0.541      output 90.3",
               "planet first :  tiredness 0.542      output 90.3",
               "profit first :  tiredness 0.542      output 90.3",
               "                    ^                    ^",
               "                 identical            identical"])
    P("Why? Measuring it gave the answer:", bold=True)
    code(doc, ["after the forbidden rules filter the options...",
               "in 95.7% of 45,754 decisions, either 0 or 1 option is left"])
    B(["An analogy — the restaurant",
       "You love chicken and hate fish. But by the time the kitchen's rules "
       "are applied, one dish is left on the menu.",
       "Do your preferences matter now? No. There is only one thing to order.",
       "When there is nothing to choose between, preferences change nothing."])
    B(["This is good news, not bad",
       "It means no change of priorities can strip the protection away. Even "
       "with profit weighted at 100%, the ban on exceeding the tiredness limit "
       "is still there.",
       "“This framework is steered by its constraints — not by its "
       "priorities.”"], fill="E9F5EC", color=GOOD)

    # ---- 9 ---------------------------------------------------------------
    H("9", "The weaknesses, stated honestly")
    P("This is what makes it real research — nothing is hidden.",
      italic=True, color=MUTED)
    T(["#", "Weakness", "Why it is written down"],
      [["1", "Posture risk fell only 13–15%, not the 20% predicted",
        "our own prediction was not met — but there is an arithmetic reason"],
       ["2", "Work was shared out less evenly",
        "a result against us — reported in full anyway"],
       ["3", "Two rates were calibrated, not measured",
        "varied from half to double; the conclusion did not change"],
       ["4", "The explanation layer was not built",
        "it is in the design, not in the code — and the paper says so"],
       ["5", "Never tested in a real factory",
        "it is a simulation, and that is stated plainly"]],
      widths=[0.35, 2.6, 3.4])
    B(["There is something good hidden in #2",
       "Asking why the work was shared unevenly gave the answer: OP1, the "
       "strongest worker, is the only one barred from heavy jobs — his skill "
       "there is 0.30 against a floor of 0.40. So heavy work falls to the two "
       "least able to sustain it.",
       "Training that one skill closes about two thirds of the gap and raises "
       "output as well.",
       "A machine-only system could not even ask this question — it has no "
       "model of the person to ask it about."])

    # ---- 10 --------------------------------------------------------------
    H("10", "The whole thing in three sentences")
    for n, line in enumerate([
            "Industry 5.0 puts the person at the centre. But digital twins are "
            "still about machines; the person is only static data.",
            "This work makes tiredness a live state, connects it to the machine's "
            "twin, and turns human limits from a “fine” into a “ban”.",
            "The result: tiredness down 27%, electricity down 31%, rules broken "
            "from 79.5 to zero — with no measurable loss of output."], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(8)
        style_run(p.add_run(f"{n}.  "), EN, 13, True, color=ACCENT)
        style_run(p.add_run(line), EN, 12)

    doc.add_paragraph()
    B(["Where to look next:",
       "paper.docx — the full paper  ·  docs/14-demo-guide.md — how to "
       "demonstrate it  ·  docs/12-paper-blueprint.md — every detail"],
      fill="F4F4F2", color=MUTED, size=10)
    return doc


def main() -> None:
    lang = (sys.argv[1] if len(sys.argv) > 1 else "en").lower()
    if lang.startswith("si"):
        doc, out = build_si(), ROOT / "doc_sinhala.docx"
    else:
        doc, out = build_en(), OUT
    doc.save(str(out))
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"  [ok] wrote {out.name}  ({lang})")
    print(f"       {len(doc.paragraphs)} paragraphs, ~{words} words, "
          f"{len(doc.tables)} tables")


if __name__ == "__main__":
    main()
